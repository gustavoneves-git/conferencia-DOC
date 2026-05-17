import json
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app import create_app
from app.database.db import get_db, init_db
from app.repositories.document_repository import DocumentRepository
from app.repositories.generated_file_repository import GeneratedFileRepository
from app.repositories.review_repository import ReviewRepository
from app.services.corrected_document_service import CorrectedDocumentService
from app.services.docx_output_service import DocxOutputService
from app.services.final_document_service import FinalDocumentService


def _app(tmp_path):
    app = create_app()
    app.config["BASE_DIR"] = tmp_path
    app.config["DATABASE_URL"] = "sqlite:///data/test.db"
    app.config["AI_MODE"] = "mock"
    for folder in ("data", "storage/corrected", "storage/final"):
        (tmp_path / folder).mkdir(parents=True, exist_ok=True)
    init_db(app)
    return app


def _seed_review(app):
    with app.app_context():
        docs = DocumentRepository()
        reviews = ReviewRepository()
        doc_id = docs.create("QUOREN - CONTRATO.pdf", "stored.pdf", str(app.config["BASE_DIR"] / "stored.pdf"), "CONTRATO")
        docs.update_metadata(doc_id, "CONTRATO_SOCIAL_CONSTITUICAO_LTDA", "QUOREN EMPREENDIMENTOS LTDA")
        docs.save_pages(
            doc_id,
            [
                {
                    "page_number": 1,
                    "text": (
                        "CONTRATO SOCIAL\n"
                        "QUOREN EMPREENDIMENTOS LTDA\n"
                        "CLÁUSULA PRIMEIRA\n"
                        "O sócio administrador, poderá representar a sociedade por mandado judicial.\n"
                        "pesa a cláusula restritiva sobre as quotas.\n"
                        "DIASSIS assina o ato.\n"
                        "Fica eleito o foro da sede.\n"
                    ),
                }
            ],
        )
        session_id = reviews.create_session(
            doc_id,
            "COMPLETED",
            {"total": 5, "BAIXA": 0, "MEDIA": 1, "ALTA": 0, "CRITICA": 0, "CONFERIR": 4},
            "mock",
            "mock",
        )
        reviews.save_issues(
            session_id,
            [
                {
                    "code": "E001",
                    "page_number": 1,
                    "original_text": "O sócio administrador, poderá",
                    "issue_type": "PONTUACAO",
                    "severity": "MEDIA",
                    "explanation": "Vírgula indevida.",
                    "technical_reason": "Não se separa sujeito e verbo por vírgula.",
                    "suggestion": "O sócio administrador poderá",
                    "recommended_action": "Corrigir pontuação.",
                    "source": "RULE",
                },
                {
                    "code": "E002",
                    "page_number": 1,
                    "original_text": "mandado judicial",
                    "issue_type": "DADO_A_CONFERIR",
                    "severity": "CONFERIR",
                    "explanation": "Validar termo.",
                    "technical_reason": "Pode haver troca entre mandado e mandato.",
                    "suggestion": "mandato judicial",
                    "recommended_action": "Conferir antes de alterar.",
                    "source": "RULE",
                },
                {
                    "code": "E003",
                    "page_number": 1,
                    "original_text": "pesa a cláusula restritiva",
                    "issue_type": "CLAUSULA_JURIDICA_SENSIVEL",
                    "severity": "CONFERIR",
                    "explanation": "Cláusula sensível.",
                    "technical_reason": "Exige validação humana.",
                    "suggestion": "incidem sobre as quotas as cláusulas restritivas de incomunicabilidade e impenhorabilidade",
                    "recommended_action": "Validar com responsável.",
                    "source": "RULE",
                },
                {
                    "code": "E004",
                    "page_number": 1,
                    "original_text": "DIASSIS",
                    "issue_type": "DADO_A_CONFERIR",
                    "severity": "CONFERIR",
                    "explanation": "Conferir grafia do nome.",
                    "technical_reason": "Nome próprio suspeito exige validação humana.",
                    "suggestion": "conferir grafia do nome",
                    "recommended_action": "Validar documento original.",
                    "source": "RULE",
                },
                {
                    "code": "E005",
                    "page_number": 1,
                    "original_text": "foro da sede",
                    "issue_type": "DADO_A_CONFERIR",
                    "severity": "CONFERIR",
                    "explanation": "Validar comarca aplicável.",
                    "technical_reason": "A comarca não pode ser inferida com segurança.",
                    "suggestion": "Fica eleito o foro da Comarca de [cidade/UF da sede]...",
                    "recommended_action": "Validar comarca.",
                    "source": "RULE",
                },
            ],
        )
        return doc_id, session_id


def test_corrected_mock_applies_only_safe_corrections(tmp_path):
    app = _app(tmp_path)
    doc_id, session_id = _seed_review(app)
    with app.app_context():
        docx, pdf = CorrectedDocumentService().generate(doc_id, session_id)
        text = "\n".join(p.text for p in DocxDocument(docx).paragraphs)
        audit = GeneratedFileRepository().latest_for_session(doc_id, session_id, "JSON_CORRECOES_APLICADAS")
        audit_data = json.loads(Path(audit["path"]).read_text(encoding="utf-8"))

    assert Path(docx).exists()
    assert Path(pdf).exists()
    assert "O sócio administrador poderá" in text
    assert "mandado judicial" in text
    assert "pesa a cláusula restritiva" in text
    assert audit_data["status_saida"] == "CORRIGIDO_PARA_REVISAO"
    assert [c["codigo"] for c in audit_data["correcoes_aplicadas"]] == ["E001"]
    assert {a["codigo"] for a in audit_data["alertas"]} == {"E002", "E003", "E004", "E005"}


def test_final_requires_human_confirmation_and_corrected_document(tmp_path):
    app = _app(tmp_path)
    doc_id, session_id = _seed_review(app)
    with app.app_context():
        with pytest.raises(ValueError):
            FinalDocumentService().generate(doc_id, session_id, False)
        with pytest.raises(ValueError):
            FinalDocumentService().generate(doc_id, session_id, True)
        CorrectedDocumentService().generate(doc_id, session_id)
        docx, pdf = FinalDocumentService().generate(doc_id, session_id, True)
        text = "\n".join(p.text for p in DocxDocument(docx).paragraphs)
        status = DocumentRepository().get(doc_id)["status"]

    assert Path(docx).exists()
    assert Path(pdf).exists()
    assert status == "FINAL_PARA_PROTOCOLO"
    assert "inteligência artificial" not in text.lower()
    assert "alerta interno" not in text.lower()


def test_docx_output_formats_title_headings_and_signature(tmp_path):
    path = tmp_path / "documento.docx"
    DocxOutputService().create_docx(
        str(path),
        "Documento corrigido para revisão",
        (
            "CONTRATO SOCIAL DE CONSTITUIÇÃO DE SOCIEDADE LIMITADA\n"
            "EMPRESA TESTE LTDA\n\n"
            "CLÁUSULA PRIMEIRA\n"
            "Texto   do contrato.\n\n\n\n"
            "São Paulo, 17 de maio de 2026.\n\n"
            "____________________________\n"
            "JOÃO DA SILVA\n"
            "Administrador\n"
        ),
        "EMPRESA TESTE LTDA",
    )
    doc = DocxDocument(path)
    section = doc.sections[0]
    assert doc.paragraphs[0].alignment == 1
    assert round(section.left_margin.cm, 1) == 3.0
    assert round(section.right_margin.cm, 1) == 2.5
    assert any(p.text == "CLÁUSULA PRIMEIRA" and p.runs[0].bold for p in doc.paragraphs if p.runs)
    assert any(p.text == "________________________________________" and p.alignment == 1 for p in doc.paragraphs)
    assert any(p.text == "JOÃO DA SILVA" and p.alignment == 1 and p.runs[0].bold for p in doc.paragraphs if p.runs)
    assert not any("  " in p.text for p in doc.paragraphs)
    assert not any(
        not doc.paragraphs[index].text.strip()
        and not doc.paragraphs[index + 1].text.strip()
        and not doc.paragraphs[index + 2].text.strip()
        for index in range(len(doc.paragraphs) - 2)
    )


def test_corrected_and_final_file_names_are_professional(tmp_path):
    app = _app(tmp_path)
    doc_id, session_id = _seed_review(app)
    with app.app_context():
        corrected_docx, corrected_pdf = CorrectedDocumentService().generate(doc_id, session_id)
        final_docx, final_pdf = FinalDocumentService().generate(doc_id, session_id, True)
    assert "__corrigido_revisao__" in Path(corrected_docx).name
    assert "__corrigido_revisao__" in Path(corrected_pdf).name
    assert "__final_protocolo__" in Path(final_docx).name
    assert "__final_protocolo__" in Path(final_pdf).name


def test_pdf_final_generates_with_structured_document(tmp_path):
    path = tmp_path / "final.pdf"
    DocxOutputService().create_pdf(
        str(path),
        "Documento final para protocolo",
        (
            "CONTRATO SOCIAL DE CONSTITUIÇÃO DE SOCIEDADE LIMITADA\n"
            "EMPRESA TESTE LTDA\n"
            "CLÁUSULA PRIMEIRA\n"
            "Texto do contrato societário.\n\n"
            "________________________________________\n"
            "MARIA DA SILVA\n"
            "Administradora"
        ),
        "EMPRESA TESTE LTDA",
    )
    assert path.exists()
    assert path.stat().st_size > 1000


def test_final_document_preserves_protected_terms(tmp_path):
    app = _app(tmp_path)
    doc_id, session_id = _seed_review(app)
    with app.app_context():
        CorrectedDocumentService().generate(doc_id, session_id)
        docx, _pdf = FinalDocumentService().generate(doc_id, session_id, True)
        text = "\n".join(p.text for p in DocxDocument(docx).paragraphs)
    assert "mandado judicial" in text
    assert "pesa a cláusula restritiva" in text
    assert "DIASSIS" in text
    assert "foro da sede" in text
    assert "mandato judicial" not in text
    assert "incidem sobre as quotas as cláusulas restritivas" not in text
