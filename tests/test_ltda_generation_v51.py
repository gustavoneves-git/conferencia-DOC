import json
from pathlib import Path

from app import create_app
from app.generation.generation_payload_builder import build_ltda_constituicao_payload
from app.services.advanced_case_detector_service import detect_advanced_case
from app.services.ltda_generation_service import LtdaGenerationService
from app.services.ltda_generation_validation_service import LtdaGenerationValidationService


def valid_form_data() -> dict:
    return {
        "razao_social": "EMPRESA FICTICIA TESTE LTDA",
        "nome_fantasia": "EMPRESA TESTE",
        "sede_logradouro": "Rua das Minutas",
        "sede_numero": "100",
        "sede_complemento": "Sala 1",
        "sede_bairro": "Centro",
        "sede_cidade": "São Paulo",
        "sede_uf": "SP",
        "sede_cep": "01000-000",
        "objeto_social": "Prestação de serviços administrativos e participação em outras sociedades.",
        "capital_social": "10000,00",
        "quotas_totais": "10000",
        "valor_unitario_quota": "1,00",
        "forma_integralizacao": "ato",
        "foro": "São Paulo/SP",
        "data_instrumento": "17 de maio de 2026",
        "cidade_assinatura": "São Paulo",
        "uf_assinatura": "SP",
        "socio1_tipo_pessoa": "PF",
        "socio1_nome": "JOAO TESTE DA SILVA",
        "socio1_nacionalidade": "brasileiro",
        "socio1_naturalidade": "São Paulo/SP",
        "socio1_estado_civil": "solteiro",
        "socio1_profissao": "empresário",
        "socio1_rg": "RG-TESTE-1",
        "socio1_orgao_expedidor": "SSP/SP",
        "socio1_cpf": "CPF-TESTE-1",
        "socio1_endereco": "Rua Um, 10, São Paulo/SP",
        "socio1_quotas": "5000",
        "socio1_percentual": "50",
        "socio2_tipo_pessoa": "PF",
        "socio2_nome": "MARIA TESTE DE SOUZA",
        "socio2_nacionalidade": "brasileira",
        "socio2_naturalidade": "São Paulo/SP",
        "socio2_estado_civil": "casada",
        "socio2_regime_bens": "comunhão parcial de bens",
        "socio2_profissao": "empresária",
        "socio2_rg": "RG-TESTE-2",
        "socio2_orgao_expedidor": "SSP/SP",
        "socio2_cpf": "CPF-TESTE-2",
        "socio2_endereco": "Rua Dois, 20, São Paulo/SP",
        "socio2_quotas": "5000",
        "socio2_percentual": "50",
        "administrador_nome": "JOAO TESTE DA SILVA",
        "administrador_genero": "masculino",
        "tipo_administracao": "isolada",
        "pro_labore": "previsto_sem_valor",
        "advogado_nome": "ADVOGADO TESTE",
        "advogado_oab": "000000",
        "advogado_oab_uf": "SP",
        "testemunha1_nome": "TESTEMUNHA UM",
        "testemunha1_documento": "DOC-1",
        "testemunha2_nome": "TESTEMUNHA DOIS",
        "testemunha2_documento": "DOC-2",
    }


def _app(tmp_path, monkeypatch):
    monkeypatch.setenv("DATABASE_URL", "sqlite:///" + str(tmp_path / "v51.db").replace("\\", "/"))
    monkeypatch.setenv("AI_MODE", "mock")
    app = create_app()
    app.config["BASE_DIR"] = tmp_path
    for folder in ("storage/generated", "data"):
        (tmp_path / folder).mkdir(parents=True, exist_ok=True)
    return app


def test_payload_builder_creates_expected_structure():
    payload = build_ltda_constituicao_payload(valid_form_data())
    assert payload["document_type"] == "CONSTITUICAO_LTDA_PADRAO"
    assert payload["empresa"]["razao_social"] == "EMPRESA FICTICIA TESTE LTDA"
    assert len(payload["socios"]) == 2
    assert payload["metadata"]["ocr_used"] is False


def test_ltda_validation_accepts_standard_case():
    payload = build_ltda_constituicao_payload(valid_form_data())
    result = LtdaGenerationValidationService().validate(payload)
    assert result["valid"] is True
    assert result["errors"] == []


def test_ltda_validation_blocks_wrong_quota_sum():
    data = valid_form_data()
    data["socio2_quotas"] = "4000"
    payload = build_ltda_constituicao_payload(data)
    result = LtdaGenerationValidationService().validate(payload)
    assert result["valid"] is False
    assert any(error["field"] == "quotas" for error in result["errors"])


def test_ltda_validation_blocks_admin_outside_socios():
    data = valid_form_data()
    data["administrador_nome"] = "ADMINISTRADOR TERCEIRO"
    payload = build_ltda_constituicao_payload(data)
    result = LtdaGenerationValidationService().validate(payload)
    assert any(error["field"] == "administrador_nome" for error in result["errors"])


def test_advanced_detector_flags_sensitive_and_complex_cases():
    data = valid_form_data()
    data["socio1_tipo_pessoa"] = "PJ"
    data["capital_em_bens"] = "on"
    data["incomunicabilidade"] = "on"
    payload = build_ltda_constituicao_payload(data)
    result = detect_advanced_case(payload)
    codes = {item["code"] for item in result["advanced_reasons"]}
    assert result["is_advanced"] is True
    assert {"SOCIO_PJ", "CAPITAL_EM_BENS", "INCOMUNICABILIDADE"} <= codes


def test_generation_mock_creates_docx_pdf_and_payload(tmp_path, monkeypatch):
    app = _app(tmp_path, monkeypatch)
    with app.app_context():
        result = LtdaGenerationService().generate_constituicao(valid_form_data())
    assert result["status"] == "ok"
    assert Path(result["files"]["docx"]).exists()
    assert Path(result["files"]["pdf"]).exists()
    assert Path(result["files"]["payload"]).exists()
    text = Path(result["files"]["payload"]).read_text(encoding="utf-8")
    payload = json.loads(text)
    assert payload["empresa"]["razao_social"] == "EMPRESA FICTICIA TESTE LTDA"
    assert "ocr_used" in payload["metadata"]


def test_generation_mock_does_not_invent_sensitive_data(tmp_path, monkeypatch):
    app = _app(tmp_path, monkeypatch)
    with app.app_context():
        result = LtdaGenerationService().generate_constituicao(valid_form_data())
    from docx import Document

    doc = Document(result["files"]["docx"])
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "CPF-TESTE-1" in text
    assert "RG-TESTE-1" in text
    assert "inteligência artificial" not in text.lower()
    assert "OpenAI" not in text
    assert "ChatGPT" not in text


def test_generated_minute_normalization_removes_initial_duplicate_title_and_company():
    service = LtdaGenerationService()
    text = """CONTRATO SOCIAL DE CONSTITUIÇÃO DE SOCIEDADE LIMITADA
EMPRESA FICTICIA TESTE LTDA

CONTRATO SOCIAL DE CONSTITUIÇÃO DE SOCIEDADE LIMITADA

Pelo presente instrumento particular, os sócios resolvem constituir sociedade.

CLÁUSULA PRIMEIRA - DA DENOMINAÇÃO
A sociedade adotará o nome empresarial EMPRESA FICTICIA TESTE LTDA.
"""
    normalized = service._normalize_generated_minute_text(
        text,
        "Contrato Social de Constituição de Sociedade Limitada",
        "EMPRESA FICTICIA TESTE LTDA",
    )
    lines = [line.strip() for line in normalized.splitlines() if line.strip()]
    assert lines[0] == "Pelo presente instrumento particular, os sócios resolvem constituir sociedade."
    assert "CLÁUSULA PRIMEIRA - DA DENOMINAÇÃO" in normalized
    assert "EMPRESA FICTICIA TESTE LTDA" in normalized


def test_generated_minute_normalization_keeps_text_without_header():
    service = LtdaGenerationService()
    text = """Pelo presente instrumento particular, os sócios resolvem constituir sociedade.

CLÁUSULA PRIMEIRA - DA DENOMINAÇÃO
A sociedade adotará o nome empresarial EMPRESA FICTICIA TESTE LTDA.
"""
    normalized = service._normalize_generated_minute_text(
        text,
        "Contrato Social de Constituição de Sociedade Limitada",
        "EMPRESA FICTICIA TESTE LTDA",
    )
    assert normalized.startswith("Pelo presente instrumento")
    assert "CLÁUSULA PRIMEIRA - DA DENOMINAÇÃO" in normalized
    assert "EMPRESA FICTICIA TESTE LTDA" in normalized


def test_generation_mock_docx_has_single_header_title_and_company(tmp_path, monkeypatch):
    app = _app(tmp_path, monkeypatch)
    with app.app_context():
        result = LtdaGenerationService().generate_constituicao(valid_form_data())
    from docx import Document

    paragraphs = [p.text.strip() for p in Document(result["files"]["docx"]).paragraphs if p.text.strip()]
    assert paragraphs.count("CONTRATO SOCIAL DE CONSTITUIÇÃO DE SOCIEDADE LIMITADA") == 1
    assert paragraphs.count("EMPRESA FICTICIA TESTE LTDA") == 1
    full_text = "\n".join(paragraphs)
    assert "CLÁUSULA PRIMEIRA - DA DENOMINAÇÃO" in full_text
    assert "CPF-TESTE-1" in full_text
    assert "RG-TESTE-1" in full_text


def test_generation_blocks_advanced_case_without_files(tmp_path, monkeypatch):
    data = valid_form_data()
    data["tipo_administracao"] = "conjunta"
    app = _app(tmp_path, monkeypatch)
    with app.app_context():
        result = LtdaGenerationService().generate_constituicao(data)
    assert result["status"] == "advanced_blocked"
    assert result["files"] == {}
    assert any(item["code"] == "ADMINISTRACAO_CONJUNTA" for item in result["advanced"]["advanced_reasons"])


def test_ltda_generation_form_route_loads(tmp_path, monkeypatch):
    app = _app(tmp_path, monkeypatch)
    response = app.test_client().get("/generation/ltda/constituicao")
    assert response.status_code == 200
    assert "Gerar minuta para revisão" in response.get_data(as_text=True)


def test_ltda_generation_post_standard_case(tmp_path, monkeypatch):
    app = _app(tmp_path, monkeypatch)
    response = app.test_client().post("/generation/ltda/constituicao/generate", data=valid_form_data())
    assert response.status_code == 200
    text = response.get_data(as_text=True)
    assert "Minuta gerada para revisão humana" in text
    assert "Baixar DOCX" in text


def test_ltda_generation_post_advanced_case(tmp_path, monkeypatch):
    data = valid_form_data()
    data["socio1_tipo_pessoa"] = "PJ"
    app = _app(tmp_path, monkeypatch)
    response = app.test_client().post("/generation/ltda/constituicao/generate", data=data)
    assert response.status_code == 200
    text = response.get_data(as_text=True)
    assert "Caso avançado bloqueado" in text
    assert "Sócio pessoa jurídica" in text


def test_existing_review_upload_route_still_works_after_v51(tmp_path, monkeypatch):
    app = _app(tmp_path, monkeypatch)
    response = app.test_client().get("/documents/upload")
    assert response.status_code == 200
