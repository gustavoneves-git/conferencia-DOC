import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer


class DocxOutputService:
    def create_docx(self, path: str, title: str, text: str, company: str | None = None) -> str:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        styles = doc.styles["Normal"]
        styles.font.name = "Times New Roman"
        styles.font.size = Pt(12)
        if title:
            self._centered(doc, title.upper(), bold=True, size=13)
        if company:
            self._centered(doc, company.upper(), bold=True, size=12)
        if title or company:
            doc.add_paragraph()

        previous_signature = False
        for line in self._clean_text(text).splitlines():
            if not line.strip():
                doc.add_paragraph()
                previous_signature = False
                continue
            p = doc.add_paragraph()
            stripped = line.strip()
            run = p.add_run(stripped)
            role = self._line_role(stripped, previous_signature)
            if role in {"title", "company"}:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
            elif role in {"clause", "article", "major_heading", "paragraph_heading"}:
                run.bold = True
                if role == "major_heading":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(8)
            elif role == "signature_line":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(3)
            elif role in {"signature_name", "signature_detail"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_together = True
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(5 if role == "signature_detail" else 2)
                if role == "signature_name":
                    run.bold = True
            elif role == "date":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(14)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(7)
            p.paragraph_format.line_spacing = 1.15
            previous_signature = role == "signature_line"
        doc.save(path)
        return path

    def create_pdf(self, path: str, title: str, text: str, company: str | None = None) -> str:
        doc = SimpleDocTemplate(
            path,
            pagesize=A4,
            rightMargin=2.5 * cm,
            leftMargin=3.0 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )
        styles = self._pdf_styles()
        story = []
        if title:
            story.append(Paragraph(self._escape(title.upper()), styles["Title"]))
        if company:
            story.append(Paragraph(self._escape(company.upper()), styles["Company"]))
        if title or company:
            story.append(Spacer(1, 18))
        lines = self._clean_text(text).splitlines()
        index = 0
        previous_signature = False
        while index < len(lines):
            raw = lines[index]
            stripped = raw.strip()
            if not stripped:
                story.append(Spacer(1, 8))
                previous_signature = False
                index += 1
                continue
            role = self._line_role(stripped, previous_signature)
            if role == "signature_line":
                block = [Paragraph(self._escape(stripped), styles["SignatureLine"])]
                lookahead = index + 1
                while lookahead < len(lines) and lines[lookahead].strip():
                    next_role = self._line_role(lines[lookahead].strip(), lookahead == index + 1)
                    if next_role not in {"signature_name", "signature_detail"}:
                        break
                    block.append(Paragraph(self._escape(lines[lookahead].strip()), styles["SignatureName" if next_role == "signature_name" else "SignatureDetail"]))
                    lookahead += 1
                story.append(KeepTogether(block))
                previous_signature = False
                index = lookahead
                continue
            if role in {"title", "company"}:
                style = styles["DocumentTitle" if role == "title" else "CompanyLine"]
            elif role in {"clause", "article", "paragraph_heading"}:
                style = styles["Heading"]
            elif role == "major_heading":
                style = styles["MajorHeading"]
            elif role in {"signature_name", "signature_detail"}:
                style = styles["SignatureName" if role == "signature_name" else "SignatureDetail"]
            elif role == "date":
                style = styles["Date"]
            else:
                style = styles["Body"]
            story.append(Paragraph(self._escape(stripped), style))
            previous_signature = role == "signature_line"
            index += 1
        doc.build(story)
        return path

    def _centered(self, doc, text, bold=False, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(6)

    def _is_heading(self, text: str) -> bool:
        return self._line_role(text) in {"title", "clause", "article", "major_heading", "paragraph_heading"}

    def _is_major_heading(self, text: str) -> bool:
        return self._line_role(text) == "major_heading"

    def _is_signature_line(self, text: str) -> bool:
        return self._line_role(text) == "signature_line"

    def _is_location_date(self, text: str) -> bool:
        return bool(re.search(r",\s*\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\.?$", text, re.IGNORECASE))

    def _line_role(self, text: str, previous_signature: bool = False) -> str:
        stripped = text.strip()
        upper = stripped.upper()
        if not stripped:
            return "blank"
        if "________________" in stripped:
            return "signature_line"
        if upper in {"ASSINATURAS", "SÓCIOS", "SOCIOS", "TESTEMUNHAS"}:
            return "major_heading"
        if previous_signature:
            if self._looks_like_signature_detail(stripped):
                return "signature_detail"
            return "signature_name"
        if self._is_location_date(stripped):
            return "date"
        if re.match(r"^(CONTRATO SOCIAL|ATA DA ASSEMBLEIA|ASSEMBLEIA GERAL|ESTATUTO SOCIAL)", upper):
            return "title"
        if re.match(r"^(CAP[ÍI]TULO|SE[CÇ][AÃ]O)\s+[IVXLCDM0-9]+", upper):
            return "major_heading"
        if upper.startswith(("CLÁUSULA", "CLAUSULA")):
            return "clause"
        if re.match(r"^(ARTIGO|ART\.)\s*\d+", upper):
            return "article"
        if upper.startswith(("PARÁGRAFO", "PARAGRAFO")):
            return "paragraph_heading"
        if self._looks_like_company_line(stripped):
            return "company"
        return "body"

    def _looks_like_company_line(self, text: str) -> bool:
        upper = text.upper()
        if len(text) > 140 or len(text.split()) < 2:
            return False
        company_terms = (" LTDA", " S/A", " S.A", " SOCIEDADE", " EIRELI", " LIMITADA")
        return upper == text and any(term in upper for term in company_terms)

    def _looks_like_signature_detail(self, text: str) -> bool:
        upper = text.upper()
        return bool(
            re.search(r"\b(RG|CPF|OAB|ADMINISTRADOR|ADMINISTRADORA|DIRETOR|DIRETORA|S[ÓO]CIO|S[ÓO]CIA|TESTEMUNHA)\b", upper)
            or len(text) <= 45 and text[:1].islower()
        )

    def _clean_text(self, text: str | None) -> str:
        lines = []
        blank = 0
        for raw in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            line = re.sub(r"[ \t]+", " ", raw).strip()
            line = re.sub(r"_{10,}", "________________________________________", line)
            if not line:
                blank += 1
                if blank <= 1:
                    lines.append("")
                continue
            blank = 0
            lines.append(line)
        return "\n".join(lines).strip()

    def _pdf_styles(self):
        base = getSampleStyleSheet()
        return {
            "Title": ParagraphStyle(
                "DocumentTitle", parent=base["Title"], fontName="Times-Bold",
                fontSize=13, leading=16, alignment=1, spaceAfter=8,
            ),
            "DocumentTitle": ParagraphStyle(
                "DocumentTitleLine", parent=base["Title"], fontName="Times-Bold",
                fontSize=13, leading=16, alignment=1, spaceBefore=6, spaceAfter=8,
                keepWithNext=True,
            ),
            "Company": ParagraphStyle(
                "Company", parent=base["Normal"], fontName="Times-Bold",
                fontSize=12, leading=15, alignment=1, spaceAfter=6,
            ),
            "CompanyLine": ParagraphStyle(
                "CompanyLine", parent=base["Normal"], fontName="Times-Bold",
                fontSize=12, leading=15, alignment=1, spaceAfter=8,
                keepWithNext=True,
            ),
            "MajorHeading": ParagraphStyle(
                "MajorHeading", parent=base["Heading2"], fontName="Times-Bold",
                fontSize=12, leading=15, alignment=1, spaceBefore=12, spaceAfter=8,
                keepWithNext=True,
            ),
            "Heading": ParagraphStyle(
                "Heading", parent=base["Normal"], fontName="Times-Bold",
                fontSize=12, leading=15, alignment=0, spaceBefore=10, spaceAfter=7,
                keepWithNext=True,
            ),
            "Body": ParagraphStyle(
                "Body", parent=base["Normal"], fontName="Times-Roman",
                fontSize=12, leading=15, alignment=4, firstLineIndent=1.25 * cm,
                spaceAfter=7, splitLongWords=False,
            ),
            "SignatureLine": ParagraphStyle(
                "SignatureLine", parent=base["Normal"], fontName="Times-Roman",
                fontSize=12, leading=16, alignment=1, spaceBefore=24, spaceAfter=2,
                keepTogether=True,
            ),
            "SignatureName": ParagraphStyle(
                "SignatureName", parent=base["Normal"], fontName="Times-Bold",
                fontSize=12, leading=15, alignment=1, spaceBefore=0, spaceAfter=2,
                keepTogether=True,
            ),
            "SignatureDetail": ParagraphStyle(
                "SignatureDetail", parent=base["Normal"], fontName="Times-Roman",
                fontSize=11, leading=14, alignment=1, spaceBefore=0, spaceAfter=5,
                keepTogether=True,
            ),
            "Date": ParagraphStyle(
                "Date", parent=base["Normal"], fontName="Times-Roman",
                fontSize=12, leading=15, alignment=2, spaceBefore=14, spaceAfter=8,
            ),
        }

    def _escape(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
